"""
Convert Manuscript_EN_v1_draft.md to formatted .docx for BBA-MBD submission.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

INPUT_MD  = r"D:\R_projects\revision_analysis\Manuscript_EN_v1_draft.md"
OUTPUT_DOCX = r"D:\R_projects\revision_analysis\Manuscript_EN_v1_draft.docx"

# ── helpers ──────────────────────────────────────────────────

def add_heading_paragraph(doc, text, font_size=11, bold=True, space_before=12, space_after=6):
    """Add a heading-style paragraph (not using built-in heading styles for cleaner control)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    return p

def add_body_paragraph(doc, text, space_after=6, first_line_indent=None, alignment=None):
    """Add a body text paragraph. Handles inline **bold** and `code` markers."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if alignment is not None:
        p.alignment = alignment

    # Parse inline formatting tokens: **bold**, `code`, and ***bold italic***
    # Split on tokens while keeping them
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = p.add_run(part[3:-3])
            run.font.bold = True
            run.font.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    return p

def remove_figure_placeholders(lines):
    """Remove ASCII-art figure placeholder blocks (═══ boxes), keep the Figure caption line."""
    out = []
    in_block = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('╔') or stripped.startswith('║') or stripped.startswith('╚'):
            in_block = True
            continue
        if in_block:
            # Check if this is the end — blank line or content after the block
            if stripped == '' or stripped == '```':
                in_block = False
                continue
            # If it's a paragraph starting with **Figure — that's the caption
            if stripped.startswith('**Figure'):
                out.append(line)
                in_block = False
                continue
            in_block = False
        # Skip empty ``` lines
        if stripped == '```':
            continue
        out.append(line)
    return out

def is_section_header(line):
    """Detect ## or ### section headers."""
    return bool(re.match(r'^#{1,4}\s', line))

def get_heading_level_and_text(line):
    m = re.match(r'^(#{1,4})\s+(.+)$', line)
    if m:
        return len(m.group(1)), m.group(2)
    return 0, line

# ── main conversion ──────────────────────────────────────────

def main():
    # Read markdown
    with open(INPUT_MD, 'r', encoding='utf-8') as f:
        raw = f.read()

    lines = raw.split('\n')

    # Remove figure placeholders
    lines = remove_figure_placeholders(lines)

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # ── Parse lines into sections ──
    # First pass: extract title (first line starting with # ), running title, abstract, etc.

    i = 0
    # Skip blank lines at start
    while i < len(lines) and lines[i].strip() == '':
        i += 1

    # Title (line starts with # )
    title_text = ''
    if i < len(lines) and lines[i].startswith('# '):
        title_text = lines[i][2:]
        i += 1

    # Check for running title
    running_title = ''
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith('**Running title**'):
            running_title = stripped.replace('**Running title**', '').strip().lstrip(':').strip()
            i += 1
        elif stripped == '---' or stripped == '':
            i += 1
        else:
            break

    # ── Add Title Page content ──
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(24)
    title_run = title_p.add_run(title_text)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'

    if running_title:
        rt_p = doc.add_paragraph()
        rt_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt_p.paragraph_format.space_after = Pt(18)
        rt_run = rt_p.add_run(f"Running title: {running_title}")
        rt_run.font.size = Pt(10)
        rt_run.font.italic = True
        rt_run.font.name = 'Times New Roman'

    # ── Process body sections ──
    # Collect remaining lines into sections based on ## / ### headers
    sections = []
    current_heading = None
    current_lines = []

    while i < len(lines):
        stripped = lines[i]
        if is_section_header(stripped):
            if current_heading is not None and current_lines:
                sections.append((current_heading, current_lines))
            level, heading = get_heading_level_and_text(stripped)
            current_heading = (level, heading)
            current_lines = []
        else:
            current_lines.append(stripped)
        i += 1

    if current_heading is not None and current_lines:
        sections.append((current_heading, current_lines))

    # ── Render sections ──
    for (level, heading), sec_lines in sections:
        # Heading
        if level == 2:
            h_font = 13
            h_before = 18
        elif level == 3:
            h_font = 12
            h_before = 14
        else:
            h_font = 11
            h_before = 12

        add_heading_paragraph(doc, heading, font_size=h_font, bold=True,
                              space_before=h_before, space_after=6)

        # Render paragraph lines
        para_text = ''
        for line in sec_lines:
            stripped = line.strip()
            if stripped == '' or stripped == '---':
                # Flush current paragraph
                if para_text.strip():
                    add_body_paragraph(doc, para_text.strip())
                    para_text = ''
                continue

            # Skip horizontal rules and empty separators
            if stripped in ('---', '---'):
                continue

            # If the line starts a new paragraph (has content and previous was ended by blank)
            if para_text:
                para_text += ' ' + stripped
            else:
                para_text = stripped

        if para_text.strip():
            add_body_paragraph(doc, para_text.strip())

    # ── Save ──
    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")

if __name__ == '__main__':
    main()
