# -*- coding: utf-8 -*-
"""Build HC submission package: cover letter + COI docx, then assemble folder."""
import subprocess, os, shutil, glob
from docx import Document
from docx.shared import Pt, Cm, RGBColor

BASE = r"D:\R_projects\revision_analysis"
OUT = os.path.join(BASE, "HC_Submission")
FONT = "Times New Roman"

def pandoc(md, docx, extra=None):
    cmd = ["pandoc", os.path.join(BASE, md), "-o", os.path.join(BASE, docx), "--from", "markdown"]
    if extra:
        cmd[4:4] = extra
    subprocess.run(cmd, check=True)

def style_simple(docx_path, size=11, spacing=1.15):
    """Minimal professional styling for letters/documents."""
    doc = Document(docx_path)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(size)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.line_spacing = spacing
    pf.space_after = Pt(6)
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = spacing
        for run in para.runs:
            run.font.name = FONT
            run.font.color.rgb = RGBColor(0, 0, 0)
    for section in doc.sections:
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    doc.save(docx_path)

def sup_fix(docx_path):
    """Post-process: convert HTML sup/sub remnants (from cover letter md) to real superscripts."""
    doc = Document(docx_path)
    for p in doc.paragraphs:
        for r in p.runs:
            t = r.text
            if t in ("1","2","3","*","#","1,3*","2#") :
                pass
    doc.save(docx_path)

# 1. Cover letter docx
pandoc("Cover_Letter_HC.md", "Cover_Letter_HC.docx")
style_simple(os.path.join(BASE, "Cover_Letter_HC.docx"))

# 2. COI docx
pandoc("Conflict_of_Interest_HC.md", "Conflict_of_Interest_HC.docx")
style_simple(os.path.join(BASE, "Conflict_of_Interest_HC.docx"))

print("Cover letter + COI docx built.")
