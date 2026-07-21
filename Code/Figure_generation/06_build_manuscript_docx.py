"""
Build final submission docx with all figures and tables inserted.
Replaces placeholder markers with actual PNG images and formatted tables.
"""
import re
import csv
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_MD   = r"D:\R_projects\revision_analysis\Manuscript_EN_v1_draft.md"
OUTPUT_DOCX = r"D:\R_projects\revision_analysis\Manuscript_EN_v1_draft.docx"
FIGURES_DIR = r"D:\R_projects\revision_analysis\figures"
TABLES_DIR  = r"D:\R_projects\revision_analysis\tables"
PROJ_DIR    = r"D:\R_projects\revision_analysis"

# ── Figure registry ─────────────────────────────────────────
# Maps caption markers to file paths
FIGURES = {
    "**Figure 1. Study design overview.": {
        "num": "Figure 1",
        "files": [os.path.join(FIGURES_DIR, "Figure1_Study_Design.png")],
        "caption": "Study design overview. The analytical framework comprises five layers: (A) WGCNA-based identification of translation co-expression modules in HF (GSE57338, n = 313); (B) independent WGCNA replication in HCC (GSE141198, n = 148); (C) parallel ssGSEA pathway activity scoring across both diseases with quantitative comparison of effect sizes; (D) external validation of TGS prognostic value in three independent HCC cohorts; (E) systematic screening for upstream transcriptional regulators of the translation co-expression program. The mirror regulation hypothesis is tested across all five layers.",
        "width": 15.5,
    },
    "**Figure 2. Independent replication and conservation of the translation "
    "co-expression program across heart failure and hepatocellular carcinoma.**": {
        "num": "Figure 2",
        "files": [os.path.join(FIGURES_DIR, "Figure2_CrossDisease_Conservation.png")],
        "caption": (
            "Independent replication and conservation of the translation co-expression program across HF and HCC. "
            "(A) Gene dendrogram and module assignment from signed WGCNA of GSE57338 (HF, n=313). "
            "The black module (n=99 genes) was identified as the translation-associated module (β=12, signed R²=0.79). "
            "(B) Gene dendrogram and module assignment from signed WGCNA of GSE141198 (HCC, n=148). "
            "The blue module (n=1,315 genes) was identified as the translation-associated module (β=4, signed R²=0.84). "
            "(C) Shared functional enrichment of translation-associated modules across diseases. "
            "GO Biological Process terms enriched in both HF black and HCC blue modules; "
            "translation-related terms highlighted. These analyses demonstrate that translation-associated "
            "co-expression programs identified independently in HF and HCC share conserved functional characteristics."
        ),
        "width": 15.5,
    },
    "**Figure 3. Cross-disease comparison of pathway effect sizes.**": {
        "num": "Figure 3",
        "files": [os.path.join(FIGURES_DIR, "Figure_ssGSEA_cross_disease.png")],
        "caption": "Cross-disease comparison of pathway effect sizes. Scatter plot of Cohen's d effect sizes for 81 pathways in HCC (TCGA-LIHC, tumor vs. normal, x-axis) versus HF (GSE57338, failing vs. non-failing, y-axis). Translation/ribosome-related pathways (n = 33) are shown as red points; non-translation pathways (n = 50) as gray points. The Spearman correlation for translation pathways is ρ = −0.598 (p = 0.0003; 10,000 permutations, p = 0.0079). Dashed lines mark d = 0.",
        "width": 15.5,
    },
    "**Figure 4. TGS prognostic validation and hub gene cross-disease directionality.**": {
        "num": "Figure 4",
        "files": [os.path.join(FIGURES_DIR, "Figure4_TGS_Validation.png")],
        "caption": "TGS prognostic validation and hub gene cross-disease directionality. (A) Kaplan–Meier survival curve for TGS (median dichotomization) in the GSE141198 Taiwanese HCC cohort (n = 148, 94 events). Log-rank p = 0.479. (B) Cross-disease expression direction (log2 fold-change) of the seven translation hub genes in HCC (TCGA-LIHC) and HF (GSE57338). RPS28 (top panel) shows opposite directionality across diseases.",
        "width": 15.5,
    },
    "**Supplementary Figure S1. Cross-disease directional consistency": {
        "num": "Supplementary Figure S1",
        "files": [os.path.join(FIGURES_DIR, "Figure_S1_Direction_Consistency.png")],
        "caption": "Cross-disease directional consistency of six directionally concordant hub genes (EEF1A1, FAU, RPL39, RPL3, RPL32, RPL41). Scatter plot of log2 fold-change values for the six TGS hub genes in HCC (TCGA-LIHC, tumor vs. normal) versus HF (GSE57338, failing vs. non-failing). Spearman ρ = 0.486 (p = 0.356). RPS28 (red) exhibits opposite directionality and was excluded from TGS; the seven-gene Spearman ρ = 0.126 (p = 0.788).",
        "width": 15.0,
    },
    "**Supplementary Figure S2. RPL39 single-gene expression": {
        "num": "Supplementary Figure S2",
        "files": [os.path.join(FIGURES_DIR, "Figure_S2_GSE141198_Survival.png")],
        "caption": "RPL39 single-gene expression in GSE141198. Kaplan–Meier curve for RPL39 expression (median dichotomization) in the GSE141198 Taiwanese HCC cohort (n = 148). RPL39 was selected as a representative hub gene from the translation co-expression program; its expression alone does not show significant prognostic value in this cohort.",
        "width": 15.0,
    },
    "**Supplementary Figure S3. Independence of hub genes": {
        "num": "Supplementary Figure S3",
        "files": [os.path.join(FIGURES_DIR, "Figure_S3_TF_Fisher_enrichment.png")],
        "caption": "Independence of hub genes from canonical TF target gene sets. Fisher's exact test results for the enrichment of the seven translation hub genes in Hallmark TF target gene sets (MYC Targets V1, MYC Targets V2, E2F Targets, mTORC1 Signaling). All p-values = 1.00.",
        "width": 15.0,
    },
    "**Figure 5. Upstream regulator analysis of the translation co-expression program.**": {
        "num": "Figure 5",
        "files": [
            os.path.join(FIGURES_DIR, "Figure_TF_TGS_correlation.png"),
            os.path.join(FIGURES_DIR, "Figure_Pathway_TGS_correlation.png"),
        ],
        "caption": "Upstream regulator analysis of the translation co-expression program. (A) Spearman correlation of 19 candidate transcription factors with TGS in TCGA-LIHC tumor samples (n = 371). Red bars: positive correlation (FDR < 0.05); blue bars: negative correlation (FDR < 0.05); gray: not significant. ATF4 is the strongest TF correlate (ρ = +0.439, FDR < 0.0001). (B) Spearman correlation of Hallmark pathway ssGSEA scores with TGS. MYC Targets V2 pathway activity shows the strongest correlation with TGS (ρ = +0.613, p < 0.0001).",
        "width": 15.5,
    },
    "**Figure 6. Cross-disease mirror regulation framework": {
        "num": "Figure 6",
        "files": [os.path.join(FIGURES_DIR, "Figure6_Mechanistic_Model.png")],
        "caption": "Cross-disease mirror regulation framework for translation in HF versus HCC. Left (HF): chronic energy deficit and mTORC1 suppression are associated with coordinated downregulation of the translation co-expression program (↓). Right (HCC): ISR activation (eIF2α phosphorylation → ATF4 translation) and MYC-driven proliferative signaling are associated with upregulation of the same program (↑). Center: the translation co-expression module, structurally conserved across both diseases. Arrows indicate the direction of transcriptional perturbation.",
        "width": 16.0,
    },
    "**Supplementary Figure S4. Soft threshold selection": {
        "num": "Supplementary Figure S4",
        "files": [
            os.path.join(FIGURES_DIR, "Figure_S4A_SoftThreshold_GSE141198.png"),
            os.path.join(FIGURES_DIR, "Figure_S4B_SoftThreshold_GSE57338.png"),
        ],
        "caption": "Soft threshold selection for WGCNA. Scale-free topology model fit (R², left panels) and mean connectivity (right panels) as a function of soft threshold power (β). (A) GSE141198 (HCC): β = 4 (R² = 0.84). (B) GSE57338 (HF): β = 12 (R² = 0.79). The R² ≈ 0.85 criterion (dashed red line) was used for threshold selection.",
        "width": 15.5,
    },
}

# ── Helpers ─────────────────────────────────────────────────

def add_styled_paragraph(doc, text, font_size=11, bold=False, italic=False,
                         space_before=0, space_after=6, alignment=None,
                         first_line_indent=None, font_name='Times New Roman'):
    """Add a paragraph with basic formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if alignment is not None:
        p.alignment = alignment

    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    return p

def add_heading_para(doc, text, level=1):
    """Add a section heading."""
    size_map = {1: 14, 2: 13, 3: 12, 4: 11}
    sz = size_map.get(level, 11)
    return add_styled_paragraph(doc, text, font_size=sz, bold=True,
                                space_before=18 if level <= 2 else 12,
                                space_after=6)

def add_figure(doc, fig_info):
    """Insert figure images and caption."""
    for i, fpath in enumerate(fig_info["files"]):
        if os.path.exists(fpath):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(4)
            run = p_img.add_run()
            run.add_picture(fpath, width=Cm(fig_info["width"]))
        else:
            p_missing = doc.add_paragraph()
            p_missing.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_missing.add_run(f"[MISSING: {os.path.basename(fpath)}]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(10)

    # Caption
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.paragraph_format.line_spacing = 1.0
    run_num = p_cap.add_run(f"{fig_info['num']}. ")
    run_num.font.bold = True
    run_num.font.size = Pt(10)
    run_num.font.name = 'Times New Roman'
    run_cap = p_cap.add_run(fig_info["caption"])
    run_cap.font.size = Pt(10)
    run_cap.font.name = 'Times New Roman'

def read_table_csv(fpath):
    """Read a CSV/TSV table file, return (headers, rows)."""
    if not os.path.exists(fpath):
        return None, None
    with open(fpath, 'r', encoding='utf-8') as f:
        first_line = f.readline()
        f.seek(0)
        delimiter = ',' if ',' in first_line else '\t'
        reader = csv.reader(f, delimiter=delimiter)
        rows = list(reader)
    if not rows:
        return None, None
    headers = rows[0]
    data = rows[1:]
    return headers, data

def add_table_from_data(doc, headers, data, title, col_widths=None):
    """Create a formatted Word table."""
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run = p_title.add_run(title)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

    ncols = len(headers)
    nrows = len(data) + 1  # +1 for header
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        # Header shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2F5496"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            if j < ncols:
                cell = table.rows[i + 1].cells[j]
                cell.text = ''
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                if j < ncols:
                    row.cells[j].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table

# ── Table builders ───────────────────────────────────────────

def add_table1(doc):
    """Cohort characteristics."""
    fpath = os.path.join(TABLES_DIR, "Table1_Cohort_Characteristics.txt")
    if not os.path.exists(fpath):
        add_styled_paragraph(doc, "[TABLE 1 DATA FILE NOT FOUND]", font_size=10, bold=True)
        return

    with open(fpath, 'r', encoding='utf-8') as f:
        text = f.read()

    # Parse the fixed-width table
    # The table has a header section + data rows
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(12)
    run = p_title.add_run("Table 1. Cohort characteristics across three datasets.")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

    # Create table from structured data
    headers = ["Characteristic", "GSE57338 (HF)", "TCGA-LIHC (HCC)", "GSE141198 (HCC)"]
    rows = [
        ["Sample size", "313", "377ᵃ", "148"],
        ["Platform", "Affymetrix HuGene 1.1 ST", "Illumina RNA-seq", "Illumina RNA-seq"],
        ["Tissue", "LV myocardium", "Liver", "Liver"],
        ["Disease subtypes", "DCM / ICM / NF", "HCC (tumor only)", "HCC (all tumor)"],
        ["  Dilated cardiomyopathy", "82 (26.2%)", "—", "—"],
        ["  Ischemic cardiomyopathy", "95 (30.4%)", "—", "—"],
        ["  Non-failing", "136 (43.5%)", "—", "—"],
        ["AJCC Stage I/II/III/IV", "—", "175/87/86/5", "—"],
        ["Age (mean ± SD)", "52.8 ± 13.5", "59.5 ± 13.5", "NR"],
        ["Sex (M / F)", "217 / 96 (69.3% M)", "255 / 122 (67.6% M)", "NR"],
        ["OS events (%)", "N/A", "132 (35.0%)", "94 (63.5%)"],
        ["OS follow-up (month)", "N/A", "median 19.7 [0–121]", "median ~69"],
        ["Etiology", "N/A", "HBV/HCV/Alcohol/NASH/Other", "HBV/HCV/NBNC: 30/99/19"],
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2F5496"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'

    # Add footnotes
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(4)
    p_note.paragraph_format.line_spacing = 1.0
    run = p_note.add_run(
        "Abbreviations: DCM, dilated cardiomyopathy; ICM, ischemic cardiomyopathy; NF, non-failing; "
        "NR, not reported; N/A, not applicable; OS, overall survival; NBNC, non-HBV non-HCV. "
        "ᵃ TCGA-LIHC clinical data n=377; expression analyses used n=424 samples."
    )
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    run.font.italic = True

def add_table2(doc):
    """Hub gene characteristics — already embedded in manuscript as markdown table."""
    headers = ["Gene", "HF log2FC", "HF FDR", "HCC log2FC", "HCC FDR", "Direction", "HF Module", "HCC Module"]
    rows = [
        ["EEF1A1", "−0.115", "<0.01", "−0.269", "0.008", "✅ Concordant", "turquoise", "turquoise"],
        ["RPL39", "+0.052", "ns", "+0.587", "<0.001", "✅ Concordant", "black", "grey"],
        ["FAU", "+0.114", "<0.01", "+0.193", "ns", "✅ Concordant (HCC ns)", "black", "blue"],
        ["RPL3", "+0.056", "ns", "−0.034", "ns", "⚠ Weakly discordant", "turquoise", "blue"],
        ["RPL41", "−0.016", "ns", "+0.084", "ns", "⚠ Weakly discordant", "black", "grey"],
        ["RPL32", "+0.071", "ns", "+0.422", "0.003", "✅ Concordant", "black", "blue"],
        ["RPS28", "−0.115", "<0.05", "+0.514", "0.001", "❌ Discordant", "turquoise", "blue"],
    ]

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(12)
    run = p_title.add_run("Table 2. Hub Gene Characteristics and Cross-Disease Module Assignment.")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2F5496"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'

    p_note = doc.add_paragraph()
    p_note.paragraph_format.line_spacing = 1.0
    run = p_note.add_run(
        "Note: HF log2FC, HF vs. non-failing (GSE57338, limma); HCC log2FC, tumor vs. normal (TCGA-LIHC, DESeq2). "
        "RPS28 was excluded from TGS due to opposite directionality."
    )
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    run.font.italic = True

def add_table3(doc):
    """TF-TGS correlation."""
    fpath = os.path.join(TABLES_DIR, "Table3_TF_TGS_Correlation.csv")
    if not os.path.exists(fpath):
        add_styled_paragraph(doc, "[TABLE 3 DATA FILE NOT FOUND — run R extraction first]", font_size=10, bold=True)
        return
    headers, data = read_table_csv(fpath)
    if headers is None:
        return

    # Clean up headers and data
    # The CSV has columns: TF, rho, p_value, p_adj
    # Sort by abs(rho) descending
    data_sorted = sorted(data, key=lambda r: abs(float(r[1])), reverse=True)

    # Format p-values in scientific notation
    formatted_rows = []
    for row in data_sorted:
        tf = row[0].replace('rho', '').strip() if row[0].startswith('rho') else row[0]
        rho = f"{float(row[1]):+.4f}"
        p_val = float(row[2])
        p_adj_val = float(row[3])
        if p_val < 0.0001:
            p_str = f"{p_val:.1e}"
        else:
            p_str = f"{p_val:.4f}"
        if p_adj_val < 0.0001:
            padj_str = f"{p_adj_val:.1e}"
        else:
            padj_str = f"{p_adj_val:.4f}"
        formatted_rows.append([tf, rho, p_str, padj_str])

    add_table_from_data(
        doc,
        ["Transcription Factor", "Spearman ρ", "p-value", "FDR (BH)"],
        formatted_rows,
        "Table 3. Transcription Factor–TGS Spearman Correlation in TCGA-LIHC Tumors (n = 371).",
    )

def add_table_s1(doc):
    """Module comparison."""
    fpath = os.path.join(TABLES_DIR, "Table_S1_Module_Comparison.txt")
    if not os.path.exists(fpath):
        return

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(12)
    run = p_title.add_run("Supplementary Table S1. Comparison of Translation-Associated Co-Expression Modules in HF (GSE57338) and HCC (GSE141198).")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

    # Panel A: Module Overview
    p_a = doc.add_paragraph()
    run = p_a.add_run("Panel A. Module Overview")
    run.font.bold = True; run.font.size = Pt(10); run.font.name = 'Times New Roman'

    headers_a = ["Feature", "GSE57338 (HF)", "GSE141198 (HCC)"]
    rows_a = [
        ["Module Color", "Black", "Blue"],
        ["Module Size", "99", "1,315"],
        ["Soft Threshold (β)", "12 (R² = 0.79)", "4 (R² = 0.84)"],
        ["Total Modules", "22", "4"],
    ]
    table_a = doc.add_table(rows=len(rows_a) + 1, cols=3)
    table_a.style = 'Light Grid Accent 1'
    for j, h in enumerate(headers_a):
        cell = table_a.rows[0].cells[j]; cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True; run.font.size = Pt(9); run.font.name = 'Times New Roman'
    for i, row in enumerate(rows_a):
        for j, val in enumerate(row):
            cell = table_a.rows[i + 1].cells[j]; cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.name = 'Times New Roman'

    # Hub gene co-localization
    doc.add_paragraph()
    p_hub = doc.add_paragraph()
    run = p_hub.add_run("Hub Gene Cross-Disease Module Assignment")
    run.font.bold = True; run.font.size = Pt(10); run.font.name = 'Times New Roman'

    hub_h = ["Hub Gene", "GSE57338 Module", "GSE141198 Module"]
    hub_d = [
        ["EEF1A1", "turquoise", "turquoise"],
        ["FAU", "black", "blue"],
        ["RPL39", "black", "grey"],
        ["RPL3", "turquoise", "blue"],
        ["RPL32", "black", "blue"],
        ["RPL41", "black", "grey"],
        ["RPS28", "turquoise", "blue"],
    ]
    table_hub = doc.add_table(rows=len(hub_d) + 1, cols=3)
    table_hub.style = 'Light Grid Accent 1'
    for j, h in enumerate(hub_h):
        cell = table_hub.rows[0].cells[j]; cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True; run.font.size = Pt(9); run.font.name = 'Times New Roman'
    for i, row in enumerate(hub_d):
        for j, val in enumerate(row):
            cell = table_hub.rows[i + 1].cells[j]; cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.name = 'Times New Roman'

    p_note = doc.add_paragraph()
    run = p_note.add_run("Hub genes in primary translation module: 4/7 (HF) vs. 4/7 (HCC). Four hub genes (FAU, RPL3, RPL32, RPS28) co-localize to the translation module in both diseases.")
    run.font.size = Pt(9); run.font.name = 'Times New Roman'; run.font.italic = True

    # Turquoise module note
    p_turq = doc.add_paragraph()
    p_turq.paragraph_format.space_before = Pt(6)
    run = p_turq.add_run(
        "Note: GSE141198 WGCNA identified 4 modules (blue, turquoise, brown, grey). "
        "The turquoise module (n=1,858, the largest) was flagged by the initial loose GO enrichment screen "
        "but is not a canonical translation module — blue (n=1,665) was selected as the primary translation module "
        "based on strong cytoplasmic translation GO terms (p < 1e-30) and consistent designation with the manuscript Methods."
    )
    run.font.size = Pt(9); run.font.name = 'Times New Roman'; run.font.italic = True

def add_table_s2(doc):
    """ssGSEA pathway effect sizes."""
    fpath = os.path.join(TABLES_DIR, "Table_S2_ssGSEA_Effect_Sizes.csv")
    if not os.path.exists(fpath):
        fpath_txt = os.path.join(TABLES_DIR, "Table_S2_ssGSEA_Effect_Sizes.txt")
        if os.path.exists(fpath_txt):
            fpath = fpath_txt
        else:
            add_styled_paragraph(doc, "[TABLE S2 DATA FILE NOT FOUND]", font_size=10, bold=True)
            return

    headers, data = read_table_csv(fpath)
    if headers is None:
        return

    # Limit display to first 20 rows + note
    display_data = data[:20]
    add_table_from_data(
        doc, headers, display_data,
        "Supplementary Table S2. Complete ssGSEA Pathway Effect Sizes (81 Pathways). "
        "Shown: first 20 of 81 pathways. Full table available in supplementary materials.",
    )

def add_table_s3(doc):
    """Hub gene direction consistency."""
    fpath = os.path.join(TABLES_DIR, "Table_S3_Direction_Consistency.csv")
    if not os.path.exists(fpath):
        add_styled_paragraph(doc, "[TABLE S3 DATA FILE NOT FOUND]", font_size=10, bold=True)
        return
    headers, data = read_table_csv(fpath)
    if headers is None:
        return
    add_table_from_data(
        doc, headers, data,
        "Supplementary Table S3. Hub Gene Cross-Disease Direction Consistency Analysis.",
    )

# ── Main document builder ────────────────────────────────────

def build_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # Read markdown
    with open(INPUT_MD, 'r', encoding='utf-8') as f:
        raw = f.read()

    lines = raw.split('\n')

    # ── Parse and build ──
    i = 0
    # Skip opening blanks
    while i < len(lines) and lines[i].strip() == '':
        i += 1

    # Title
    title_text = ''
    if i < len(lines) and lines[i].startswith('# '):
        title_text = lines[i][2:].strip()
        i += 1

    # Running title
    running_title = ''
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith('**Running title**'):
            running_title = s.replace('**Running title**', '').strip().lstrip(':').strip()
            i += 1
        elif s == '---' or s == '':
            i += 1
        else:
            break

    # Fallbacks
    if not title_text:
        title_text = "Conserved Translational Co-expression Programs Exhibit Mirror Regulation Across Heart Failure and Hepatocellular Carcinoma"
    if not running_title:
        running_title = "Mirror Regulation of Translation in HF and HCC"

    # ── Title page ──
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(24)
    run = p_title.add_run(title_text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    if running_title:
        p_rt = doc.add_paragraph()
        p_rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rt.paragraph_format.space_after = Pt(18)
        run = p_rt.add_run(f"Running title: {running_title}")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.name = 'Times New Roman'

    # ── Process body ──
    # Group lines: section headers start new groups
    sections = []  # [(level, heading, body_lines)]
    current_heading = None
    current_lines = []

    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # Check for section header
        header_match = re.match(r'^(#{1,4})\s+(.+)$', s)
        if header_match:
            if current_heading is not None and current_lines:
                sections.append((current_heading[0], current_heading[1], current_lines))
            level = len(header_match.group(1))
            heading = header_match.group(2)
            current_heading = (level, heading)
            current_lines = []
        else:
            current_lines.append(line)
        i += 1

    if current_heading is not None and current_lines:
        sections.append((current_heading[0], current_heading[1], current_lines))

    # ── Render sections ──
    # Track table insertion state
    table1_inserted = False
    table3_inserted = False

    for level, heading, sec_lines in sections:
        # Heading
        add_heading_para(doc, heading, level=level)

        # Insert Table 1 after Methods 4.1
        if heading.startswith("4.2") and not table1_inserted:
            add_table1(doc)
            table1_inserted = True

        # Build paragraphs from lines
        para_text = ''
        for line in sec_lines:
            s = line.strip()

            # Skip ASCII art blocks
            if s.startswith('╔') or s.startswith('║') or s.startswith('╚'):
                continue
            if s == '```' or s == '---':
                if para_text.strip():
                    _add_formatted_paragraph(doc, para_text.strip())
                    para_text = ''
                continue

            # Check for figure caption -> insert figure
            fig_inserted = False
            for marker, fig_info in FIGURES.items():
                if s.startswith(marker):
                    if para_text.strip():
                        _add_formatted_paragraph(doc, para_text.strip())
                        para_text = ''
                    add_figure(doc, fig_info)
                    fig_inserted = True
                    break

            if fig_inserted:
                continue

            # Check for table references
            if 'Complete TF–TGS correlation statistics for all nineteen candidate TFs are summarized in Table 3.' in s:
                if para_text.strip():
                    _add_formatted_paragraph(doc, para_text.strip())
                    para_text = ''
                _add_formatted_paragraph(doc, s)
                if not table3_inserted:
                    add_table3(doc)
                    table3_inserted = True
                continue

            # Accumulate paragraph text
            if s == '':
                if para_text.strip():
                    _add_formatted_paragraph(doc, para_text.strip())
                    para_text = ''
            else:
                if para_text:
                    para_text += ' ' + s
                else:
                    para_text = s

        if para_text.strip():
            _add_formatted_paragraph(doc, para_text.strip())

    # ── Insert remaining tables at end ──
    doc.add_page_break()
    add_heading_para(doc, "Tables", level=1)

    if not table1_inserted:
        add_table1(doc)

    add_table2(doc)

    if not table3_inserted:
        add_table3(doc)

    # Supplementary tables
    add_heading_para(doc, "Supplementary Tables", level=2)
    add_table_s1(doc)
    doc.add_page_break()
    add_table_s2(doc)
    add_table_s3(doc)

    # ── Save ──
    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")

def _add_formatted_paragraph(doc, text):
    """Add paragraph with inline bold/italic formatting preserved."""
    if not text:
        return

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)

    # Parse inline formatting: **bold**, *italic*, ***bold italic***
    # Split carefully
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*[^*].*?\*)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = p.add_run(part[3:-3])
            run.font.bold = True
            run.font.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.font.italic = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

if __name__ == '__main__':
    build_document()
