# =============================================================================
# Format docx: Times New Roman, body 小四(12pt), headings auto-incremented
# Usage: python format_docx.py
# =============================================================================

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Chinese sizes: 小四=12pt, 四号=14pt, 小三=15pt, 三号=16pt
FONT_NAME = "Times New Roman"
FONT_NAME_EAST = "宋体"

BODY_SIZE = Pt(12)        # 小四
H1_SIZE = Pt(16)          # 三号 (title)
H2_SIZE = Pt(14)          # 四号 (section headers)
H3_SIZE = Pt(13)          # subsection headers

STYLE_MAP = {
    "Normal":    (BODY_SIZE, False),
    "Heading 1": (H1_SIZE, True),
    "Heading 2": (H2_SIZE, True),
    "Heading 3": (H3_SIZE, True),
    "Author":    (BODY_SIZE, False),
    "Abstract":  (BODY_SIZE, False),
}


def format_docx(path):
    doc = Document(path)

    # --- Modify document-level styles ---
    for style_name, (size, bold) in STYLE_MAP.items():
        style = doc.styles[style_name]
        font = style.font
        font.name = FONT_NAME
        font.size = size
        font.bold = bold
        rpr = style.element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_NAME_EAST)
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15

    # --- Fix all paragraph runs ---
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = FONT_NAME
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = rpr.makeelement(qn('w:rFonts'), {})
                rpr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), FONT_NAME_EAST)

            if para.style.name.startswith('Heading 1'):
                run.font.size = H1_SIZE
                run.font.bold = True
            elif para.style.name.startswith('Heading 2'):
                run.font.size = H2_SIZE
                run.font.bold = True
            elif para.style.name.startswith('Heading 3'):
                run.font.size = H3_SIZE
                run.font.bold = True
            else:
                run.font.size = BODY_SIZE

    # --- Figure captions ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('Fig.') or text.startswith('Figure'):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.name = FONT_NAME
                rpr = run._element.get_or_add_rPr()
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = rpr.makeelement(qn('w:rFonts'), {})
                    rpr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), FONT_NAME_EAST)
                run.font.size = Pt(10)

    # --- Tables ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(9)

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    doc.save(path)
    print(f"  {path}")
    print(f"    Body={BODY_SIZE.pt:.0f}pt, H1={H1_SIZE.pt:.0f}pt, H2={H2_SIZE.pt:.0f}pt, H3={H3_SIZE.pt:.0f}pt, Caption=10pt")


if __name__ == '__main__':
    for p in [
        r"D:\R_projects\revision_analysis\Manuscript_EN_v1_draft.docx",
        r"D:\R_projects\revision_analysis\Manuscript_SI.docx",
    ]:
        format_docx(p)
    print("\nDone.")
